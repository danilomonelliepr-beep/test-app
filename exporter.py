import base64
import html
import io
import re
from typing import Any, Dict, List, Optional
import requests

import streamlit as st

# ReportLab (Engine PDF)
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import (
    Image,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

# Python-Docx (Engine Word)
import docx
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor


# =============================================================================
# HELPER: CONVERSIONE MERMAID IN IMMAGINE PNG
# =============================================================================


def get_mermaid_image_bytes(
    mermaid_code: str, diag_title: str = "Diagram"
) -> Optional[bytes]:
    if not mermaid_code or not str(mermaid_code).strip():
        return None

    try:
        clean_code = str(mermaid_code).strip()

        # 1. Rimuove blocchi markdown ```mermaid o ```
        clean_code = re.sub(
            r"^```(?:mermaid)?", "", clean_code, flags=re.MULTILINE
        )
        clean_code = re.sub(r"^```$", "", clean_code, flags=re.MULTILINE).strip()

        # 2. Rimuove entità HTML
        clean_code = html.unescape(clean_code)

        if not clean_code:
            return None

        # 3. Base64 URL-Safe Encoding
        graph_bytes = clean_code.encode("utf-8")
        base64_string = base64.urlsafe_b64encode(graph_bytes).decode("utf-8")

        # 4. Richiesta HTTP a mermaid.ink
        url = f"https://mermaid.ink/img/{base64_string}"
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
        }

        response = requests.get(url, headers=headers, timeout=10)

        if response.status_code == 200 and len(response.content) > 100:
            return response.content
        else:
            print(
                f"[Mermaid Error - {diag_title}] HTTP {response.status_code}: {response.text[:200]}"
            )

    except Exception as e:
        print(f"[Mermaid Exception - {diag_title}] {str(e)}")

    return None


# =============================================================================
# HELPERS UTILITY & SEARCH PROFONDA
# =============================================================================


def deep_search_key(data: Any, target_keys: List[str]) -> Optional[str]:
    """Scansiona ricorsivamente tutto il JSON per trovare la prima chiave corrispondente valida."""
    if isinstance(data, dict):
        # 1. Cerca al livello corrente
        for k, v in data.items():
            if k.lower() in [tk.lower() for tk in target_keys]:
                if v is not None and str(v).strip() and str(v).strip().upper() != "N/A":
                    if isinstance(v, (dict, list)):
                        return str(v)
                    return str(v).strip()
        # 2. Ricerca nei sotto-oggetti
        for k, v in data.items():
            res = deep_search_key(v, target_keys)
            if res:
                return res
    elif isinstance(data, list):
        for item in data:
            res = deep_search_key(item, target_keys)
            if res:
                return res
    return None


def extract_field(data: Dict[str, Any], keys: List[str], default: str = "N/A") -> str:
    """Estrae un campo cercando prima al livello radice e poi in profondità."""
    res = deep_search_key(data, keys)
    return res if res else default


def set_cell_background(cell, hex_color: str):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=80, right=80):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f"</w:tcMar>"
    )
    tcPr.append(tcMar)


# =============================================================================
# 1. GENERATORE REPORT PDF (LANDSCAPE / TUTTE LE COLONNE / DIAGRAMMI ROBUSTI)
# =============================================================================


def generate_pdf_report(
    analysis_result: Dict[str, Any],
    metadata: Dict[str, Any],
    provider: str,
    model_name: str,
) -> bytes:
    buffer = io.BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=landscape(A4),
        rightMargin=30,
        leftMargin=30,
        topMargin=30,
        bottomMargin=30,
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "DocTitle",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        textColor=colors.HexColor("#1A365D"),
        spaceAfter=4,
    )
    sub_style = ParagraphStyle(
        "DocSub",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=8.5,
        leading=11,
        textColor=colors.HexColor("#718096"),
        spaceAfter=10,
    )
    h2_style = ParagraphStyle(
        "SectionHeader",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=14,
        textColor=colors.HexColor("#1A365D"),
        spaceBefore=12,
        spaceAfter=6,
        keepWithNext=True,
    )
    body_style = ParagraphStyle(
        "BodyTextCustom",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=8,
        leading=11,
        textColor=colors.HexColor("#2D3748"),
        spaceAfter=6,
    )
    cell_style = ParagraphStyle(
        "TableCell",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=6.5,
        leading=8.5,
        textColor=colors.HexColor("#2D3748"),
    )
    header_cell_style = ParagraphStyle(
        "HeaderTableCell",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=6.5,
        leading=8.5,
        textColor=colors.white,
    )
    code_style = ParagraphStyle(
        "CodeStyle",
        parent=styles["Normal"],
        fontName="Courier",
        fontSize=6,
        leading=7.5,
        textColor=colors.HexColor("#2D3748"),
    )

    story = []

    # --- HEADER ---
    story.append(
        Paragraph("Legacy Application Knowledge Extraction", title_style)
    )
    story.append(
        Paragraph(
            f"Reverse Engineering Complete Report &nbsp;|&nbsp; Provider: <b>{html.escape(str(provider))}</b> ({html.escape(str(model_name))}) &nbsp;|&nbsp; Files Analyzed: <b>{metadata.get('file_count', 0)}</b>",
            sub_style,
        )
    )

    # --- METRICHE PRINCIPALI ---
    metric_data = [
        [
            Paragraph(
                f"<b>{metadata.get('file_count', 0)}</b><br/><font size=5.5 color='#718096'>FILES</font>",
                cell_style,
            ),
            Paragraph(
                f"<b>{metadata.get('total_line_count', 0):,}</b><br/><font size=5.5 color='#718096'>LOC</font>",
                cell_style,
            ),
            Paragraph(
                f"<b>{len(analysis_result.get('components', []))}</b><br/><font size=5.5 color='#718096'>COMPONENTS</font>",
                cell_style,
            ),
            Paragraph(
                f"<b>{len(analysis_result.get('business_rules', []))}</b><br/><font size=5.5 color='#718096'>RULES</font>",
                cell_style,
            ),
            Paragraph(
                f"<b>{len(analysis_result.get('technical_risks', []))}</b><br/><font size=5.5 color='#718096'>RISKS</font>",
                cell_style,
            ),
        ]
    ]
    t_metrics = Table(metric_data, colWidths=[156, 156, 156, 156, 156])
    t_metrics.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F7FAFC")),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                (
                    "INNERGRID",
                    (0, 0),
                    (-1, -1),
                    0.5,
                    colors.HexColor("#E2E8F0"),
                ),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E0")),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.append(t_metrics)
    story.append(Spacer(1, 8))

    # --- OVERVIEW (SUMMARY, PURPOSE, NOTES) ---
    story.append(Paragraph("1. Executive Summary & Purpose", h2_style))

    exec_summary = extract_field(
        analysis_result,
        ["executive_summary", "summary", "overview", "description"],
        default="N/A",
    )
    app_purpose = extract_field(
        analysis_result,
        ["application_purpose", "purpose", "system_purpose", "goal", "objective"],
        default="N/A",
    )
    app_notes = extract_field(
        analysis_result,
        [
            "notes",
            "additional_notes",
            "remarks",
            "comments",
            "observations",
            "extra_notes",
            "note",
            "developer_notes",
        ],
        default="No additional notes provided",
    )

    story.append(
        Paragraph(
            f"<b>Executive Summary:</b> {html.escape(exec_summary)}",
            body_style,
        )
    )
    story.append(
        Paragraph(
            f"<b>Application Purpose:</b> {html.escape(app_purpose)}",
            body_style,
        )
    )
    story.append(
        Paragraph(
            f"<b>Notes:</b> {html.escape(app_notes)}",
            body_style,
        )
    )

    story.append(Spacer(1, 6))

    # --- METRICHE CUSTOM (SE PRESENTI NEL JSON) ---
    custom_metrics = analysis_result.get(
        "metrics", analysis_result.get("code_metrics", {})
    )
    if custom_metrics and isinstance(custom_metrics, dict):
        story.append(Paragraph("<b>Additional System Metrics:</b>", body_style))
        m_rows = [
            [
                Paragraph("<b>Metric Name</b>", header_cell_style),
                Paragraph("<b>Value</b>", header_cell_style),
            ]
        ]
        for mk, mv in custom_metrics.items():
            m_rows.append(
                [
                    Paragraph(
                        html.escape(str(mk).replace("_", " ").title()),
                        cell_style,
                    ),
                    Paragraph(html.escape(str(mv)), cell_style),
                ]
            )
        t_cm = Table(m_rows, colWidths=[390, 390])
        t_cm.setStyle(
            TableStyle(
                [
                    (
                        "BACKGROUND",
                        (0, 0),
                        (-1, 0),
                        colors.HexColor("#2D3748"),
                    ),
                    ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#E2E8F0")),
                    ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E0")),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ]
            )
        )
        story.append(t_cm)
        story.append(Spacer(1, 8))

    # Helper Tabelle Dinamiche PDF
    def build_auto_pdf_table(data: List[Dict[str, Any]], total_width: int = 780):
        if not data or not isinstance(data, list):
            t = Table(
                [[Paragraph("<i>No record extracted</i>", cell_style)]],
                colWidths=[total_width],
            )
            t.setStyle(
                TableStyle(
                    [
                        (
                            "BACKGROUND",
                            (0, 0),
                            (-1, -1),
                            colors.HexColor("#F7FAFC"),
                        )
                    ]
                )
            )
            return t

        keys = []
        for item in data:
            if isinstance(item, dict):
                for k in item.keys():
                    if k not in keys:
                        keys.append(k)

        if not keys:
            return Table(
                [[Paragraph("<i>No data available</i>", cell_style)]],
                colWidths=[total_width],
            )

        if "sme_approved" in keys:
            keys.remove("sme_approved")
            keys.insert(0, "sme_approved")

        headers = [k.replace("_", " ").title() for k in keys]
        header_row = [Paragraph(f"<b>{h}</b>", header_cell_style) for h in headers]
        table_rows = [header_row]

        for item in data:
            if not isinstance(item, dict):
                continue
            row_cells = []
            for k in keys:
                if k == "sme_approved":
                    val_bool = item.get("sme_approved", False)
                    val_str = (
                        "<font color='#22543D'><b>APPROVED</b></font>"
                        if val_bool
                        else "<font color='#744210'><b>PENDING</b></font>"
                    )
                else:
                    val_str = html.escape(str(item.get(k, "-")))
                    if k == "severity":
                        sev = val_str.upper()
                        c = {
                            "CRITICAL": "#742A2A",
                            "HIGH": "#744210",
                            "MEDIUM": "#2D3748",
                            "LOW": "#234E52",
                        }.get(sev, "#2D3748")
                        val_str = f"<font color='{c}'><b>{sev}</b></font>"
                row_cells.append(Paragraph(val_str, cell_style))
            table_rows.append(row_cells)

        col_width = int(total_width / len(keys))
        col_widths = [col_width] * len(keys)

        t = Table(table_rows, colWidths=col_widths)
        t.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2D3748")),
                    ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    (
                        "INNERGRID",
                        (0, 0),
                        (-1, -1),
                        0.5,
                        colors.HexColor("#E2E8F0"),
                    ),
                    ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E0")),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ("LEFTPADDING", (0, 0), (-1, -1), 2),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 2),
                ]
            )
        )
        return t

    # --- TABELLE SEZIONI ---
    story.append(Paragraph("2. Business Logic", h2_style))
    story.append(Paragraph("<b>Business Processes:</b>", body_style))
    story.append(build_auto_pdf_table(analysis_result.get("business_processes", [])))
    story.append(Spacer(1, 6))
    story.append(Paragraph("<b>Business Rules:</b>", body_style))
    story.append(build_auto_pdf_table(analysis_result.get("business_rules", [])))

    story.append(Paragraph("3. System Architecture", h2_style))
    story.append(Paragraph("<b>Components:</b>", body_style))
    story.append(build_auto_pdf_table(analysis_result.get("components", [])))
    story.append(Spacer(1, 6))
    story.append(Paragraph("<b>Dependencies:</b>", body_style))
    story.append(build_auto_pdf_table(analysis_result.get("dependencies", [])))
    story.append(Spacer(1, 6))
    story.append(Paragraph("<b>Interfaces:</b>", body_style))
    story.append(build_auto_pdf_table(analysis_result.get("interfaces", [])))
    story.append(Spacer(1, 6))
    story.append(Paragraph("<b>Application Mapping:</b>", body_style))
    story.append(
        build_auto_pdf_table(analysis_result.get("application_mapping", []))
    )

    story.append(Paragraph("4. Data Objects & Data Flows", h2_style))
    story.append(Paragraph("<b>Data Objects:</b>", body_style))
    story.append(build_auto_pdf_table(analysis_result.get("data_objects", [])))
    story.append(Spacer(1, 6))
    story.append(Paragraph("<b>Data Flows:</b>", body_style))
    story.append(build_auto_pdf_table(analysis_result.get("data_flows", [])))

    story.append(Paragraph("5. Technical Risks & Impact Analysis", h2_style))
    story.append(Paragraph("<b>Technical Risks:</b>", body_style))
    story.append(build_auto_pdf_table(analysis_result.get("technical_risks", [])))
    story.append(Spacer(1, 6))
    story.append(Paragraph("<b>Impact Analysis:</b>", body_style))
    story.append(build_auto_pdf_table(analysis_result.get("impact_analysis", [])))

    # --- DIAGRAMMI ---
    story.append(PageBreak())
    story.append(Paragraph("6. Architecture & Process Diagrams", h2_style))

    diagrams_map = [
        (
            "Process Flow Diagram",
            analysis_result.get("mermaid_process_flow")
            or analysis_result.get("process_flow_diagram")
            or analysis_result.get("process_flow"),
        ),
        (
            "Application Mapping Diagram",
            analysis_result.get("mermaid_application_map")
            or analysis_result.get("application_map_diagram")
            or analysis_result.get("app_map"),
        ),
        (
            "Data Flow Diagram",
            analysis_result.get("mermaid_data_flow")
            or analysis_result.get("data_flow_diagram")
            or analysis_result.get("data_flow"),
        ),
        (
            "Call Graph Diagram",
            analysis_result.get("mermaid_call_graph")
            or analysis_result.get("call_graph_diagram")
            or analysis_result.get("call_graph"),
        ),
    ]

    has_diagrams = False

    for diag_title, diag_code in diagrams_map:
        if diag_code and str(diag_code).strip():
            has_diagrams = True
            story.append(Paragraph(f"<b>{diag_title}</b>", body_style))

            img_bytes = get_mermaid_image_bytes(str(diag_code), diag_title)

            if img_bytes:
                img_buf = io.BytesIO(img_bytes)
                story.append(Image(img_buf, width=720, height=240))
            else:
                clean_txt = html.escape(str(diag_code).strip()).replace(
                    "\n", "<br/>"
                )
                fallback_table = Table(
                    [
                        [
                            Paragraph(
                                f"<b>Mermaid Source Code:</b><br/>{clean_txt}",
                                code_style,
                            )
                        ]
                    ],
                    colWidths=[780],
                )
                fallback_table.setStyle(
                    TableStyle(
                        [
                            (
                                "BACKGROUND",
                                (0, 0),
                                (-1, -1),
                                colors.HexColor("#EDF2F7"),
                            ),
                            (
                                "BOX",
                                (0, 0),
                                (-1, -1),
                                0.5,
                                colors.HexColor("#CBD5E0"),
                            ),
                            ("TOPPADDING", (0, 0), (-1, -1), 6),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                        ]
                    )
                )
                story.append(fallback_table)
            story.append(Spacer(1, 10))

    if not has_diagrams:
        story.append(
            Paragraph(
                "<i>No diagrams available in the analysis result.</i>", body_style
            )
        )

    # --- STATIC EVIDENCE ---
    story.append(Spacer(1, 10))
    story.append(Paragraph("7. Static Code Evidence", h2_style))
    sql_tables = metadata.get("detected_tables", [])
    if sql_tables:
        story.append(
            Paragraph(
                f"<b>SQL Tables Extracted:</b> {html.escape(', '.join(sql_tables))}",
                body_style,
            )
        )

    doc.build(story)
    return buffer.getvalue()


# =============================================================================
# 2. GENERATORE REPORT WORD (.DOCX)
# =============================================================================


def generate_docx_report(
    analysis_result: Dict[str, Any],
    metadata: Dict[str, Any],
    provider: str,
    model_name: str,
) -> bytes:
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(8)

    doc.add_heading("Legacy Application Knowledge Extraction", level=0)

    # Sub-header con info generiche
    p_sub = doc.add_paragraph()
    p_sub.add_run("Reverse Engineering Complete Report | Provider: ").bold = False
    p_sub.add_run(f"{provider} ({model_name}) | ").bold = True
    p_sub.add_run(
        f"Files Analyzed: {metadata.get('file_count', 0)}"
    ).bold = True

    def create_auto_docx_table(data: List[Dict[str, Any]]):
        if not data or not isinstance(data, list):
            doc.add_paragraph("No record extracted")
            return

        keys = []
        for item in data:
            if isinstance(item, dict):
                for k in item.keys():
                    if k not in keys:
                        keys.append(k)

        if not keys:
            doc.add_paragraph("No data available")
            return

        if "sme_approved" in keys:
            keys.remove("sme_approved")
            keys.insert(0, "sme_approved")

        headers = [k.replace("_", " ").title() for k in keys]
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        hdr_cells = table.rows[0].cells
        for i, h_text in enumerate(headers):
            hdr_cells[i].text = h_text
            set_cell_background(hdr_cells[i], "2D3748")
            p = hdr_cells[i].paragraphs[0]
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(7.5)
                run.font.color.rgb = RGBColor(255, 255, 255)

        for item in data:
            if not isinstance(item, dict):
                continue
            row_cells = table.add_row().cells
            for i, k in enumerate(keys):
                if k == "sme_approved":
                    val = (
                        "Approved"
                        if item.get("sme_approved", False)
                        else "Pending"
                    )
                else:
                    val = str(item.get(k, "-"))

                row_cells[i].text = val
                set_cell_margins(row_cells[i])
                p = row_cells[i].paragraphs[0]
                for run in p.runs:
                    run.font.size = Pt(7)
        doc.add_paragraph()

    # --- METRICHE PRINCIPALI E STATISTICHE (DOCX) ---
    doc.add_heading("Metrics & Overview Statistics", level=2)
    t_stat = doc.add_table(rows=2, cols=5)
    t_stat.style = "Table Grid"
    stat_headers = ["FILES", "LOC", "COMPONENTS", "RULES", "RISKS"]
    stat_values = [
        str(metadata.get("file_count", 0)),
        f"{metadata.get('total_line_count', 0):,}",
        str(len(analysis_result.get("components", []))),
        str(len(analysis_result.get("business_rules", []))),
        str(len(analysis_result.get("technical_risks", []))),
    ]

    for i, h in enumerate(stat_headers):
        cell_h = t_stat.rows[0].cells[i]
        cell_h.text = h
        set_cell_background(cell_h, "2D3748")
        cell_h.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell_h.paragraphs[0].runs[0].font.bold = True

        cell_v = t_stat.rows[1].cells[i]
        cell_v.text = stat_values[i]
        cell_v.paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()

    # METRICHE CUSTOM (SE PRESENTI)
    custom_metrics = analysis_result.get(
        "metrics", analysis_result.get("code_metrics", {})
    )
    if custom_metrics and isinstance(custom_metrics, dict):
        doc.add_paragraph("Additional System Metrics:").runs[0].bold = True
        table_m = doc.add_table(rows=1, cols=2)
        table_m.style = "Table Grid"
        hdr_cells = table_m.rows[0].cells
        hdr_cells[0].text = "Metric"
        hdr_cells[1].text = "Value"
        hdr_cells[0].paragraphs[0].runs[0].font.bold = True
        hdr_cells[1].paragraphs[0].runs[0].font.bold = True

        for key, value in custom_metrics.items():
            row_cells = table_m.add_row().cells
            row_cells[0].text = str(key).replace("_", " ").title()
            row_cells[1].text = str(value)

        doc.add_paragraph()

    # --- 1. EXECUTIVE SUMMARY, PURPOSE & NOTES ---
    doc.add_heading("1. Executive Summary & Purpose", level=1)

    exec_summary = extract_field(
        analysis_result,
        ["executive_summary", "summary", "overview", "description"],
        default="N/A",
    )
    app_purpose = extract_field(
        analysis_result,
        ["application_purpose", "purpose", "system_purpose", "goal", "objective"],
        default="N/A",
    )
    app_notes = extract_field(
        analysis_result,
        [
            "notes",
            "additional_notes",
            "remarks",
            "comments",
            "observations",
            "extra_notes",
            "note",
            "developer_notes",
        ],
        default="No additional notes provided",
    )

    p_exec = doc.add_paragraph()
    p_exec.add_run("Executive Summary: ").bold = True
    p_exec.add_run(exec_summary)

    p_purp = doc.add_paragraph()
    p_purp.add_run("Application Purpose: ").bold = True
    p_purp.add_run(app_purpose)

    p_notes = doc.add_paragraph()
    p_notes.add_run("Notes: ").bold = True
    p_notes.add_run(app_notes)

    doc.add_paragraph()  # Spaziatore

    # --- SEZIONI SUCCESSIVE ---
    doc.add_heading("2. Business Logic", level=1)
    doc.add_paragraph("Business Processes:").runs[0].bold = True
    create_auto_docx_table(analysis_result.get("business_processes", []))
    doc.add_paragraph("Business Rules:").runs[0].bold = True
    create_auto_docx_table(analysis_result.get("business_rules", []))

    doc.add_heading("3. System Architecture", level=1)
    doc.add_paragraph("Components:").runs[0].bold = True
    create_auto_docx_table(analysis_result.get("components", []))
    doc.add_paragraph("Dependencies:").runs[0].bold = True
    create_auto_docx_table(analysis_result.get("dependencies", []))
    doc.add_paragraph("Interfaces:").runs[0].bold = True
    create_auto_docx_table(analysis_result.get("interfaces", []))
    doc.add_paragraph("Application Mapping:").runs[0].bold = True
    create_auto_docx_table(analysis_result.get("application_mapping", []))

    doc.add_heading("4. Data Objects & Data Flows", level=1)
    doc.add_paragraph("Data Objects:").runs[0].bold = True
    create_auto_docx_table(analysis_result.get("data_objects", []))
    doc.add_paragraph("Data Flows:").runs[0].bold = True
    create_auto_docx_table(analysis_result.get("data_flows", []))

    doc.add_heading("5. Technical Risks & Impact Analysis", level=1)
    doc.add_paragraph("Technical Risks:").runs[0].bold = True
    create_auto_docx_table(analysis_result.get("technical_risks", []))
    doc.add_paragraph("Impact Analysis:").runs[0].bold = True
    create_auto_docx_table(analysis_result.get("impact_analysis", []))

    # --- DIAGRAMMI WORD ---
    doc.add_heading("6. Architecture & Process Diagrams", level=1)
    diagrams_map = [
        (
            "Process Flow Diagram",
            analysis_result.get("mermaid_process_flow")
            or analysis_result.get("process_flow_diagram")
            or analysis_result.get("process_flow"),
        ),
        (
            "Application Mapping Diagram",
            analysis_result.get("mermaid_application_map")
            or analysis_result.get("application_map_diagram")
            or analysis_result.get("app_map"),
        ),
        (
            "Data Flow Diagram",
            analysis_result.get("mermaid_data_flow")
            or analysis_result.get("data_flow_diagram")
            or analysis_result.get("data_flow"),
        ),
        (
            "Call Graph Diagram",
            analysis_result.get("mermaid_call_graph")
            or analysis_result.get("call_graph_diagram")
            or analysis_result.get("call_graph"),
        ),
    ]

    for diag_title, diag_code in diagrams_map:
        doc.add_paragraph(diag_title).runs[0].bold = True
        if diag_code and str(diag_code).strip():
            img_bytes = get_mermaid_image_bytes(str(diag_code), diag_title)
            if img_bytes:
                img_buf = io.BytesIO(img_bytes)
                doc.add_picture(img_buf, width=Inches(9.5))
            else:
                p_code = doc.add_paragraph(str(diag_code).strip())
                p_code.style.font.name = "Courier New"
                p_code.style.font.size = Pt(7)
        else:
            doc.add_paragraph("No diagram available for this section")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
